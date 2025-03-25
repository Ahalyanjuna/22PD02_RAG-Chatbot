# Document-Only RAG Chatbot Implementation
# Uses free models only - Focused on PDF, DOCX, and CSV processing

import os
import time
import uuid
import numpy as np
import torch
import PyPDF2
import docx
import csv
import pandas as pd
import faiss
from transformers import (
    AutoTokenizer, 
    AutoModel, 
    pipeline, 
    AutoModelForSeq2SeqLM,
    AutoModelForQuestionAnswering
)
from sentence_transformers import SentenceTransformer
from flask import Flask, request, jsonify, render_template
import json
import nltk
from nltk.tokenize import sent_tokenize

# Download required NLTK resources
#nltk.download('punkt')

# Initialize Flask app
app = Flask(__name__)

# Directory setup
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
KB_FOLDER = os.path.join(BASE_DIR, 'knowledge_base')
VECTOR_DB_PATH = os.path.join(BASE_DIR, 'vector_db')

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(KB_FOLDER, exist_ok=True)
os.makedirs(VECTOR_DB_PATH, exist_ok=True)

# Load models
# Text embedding model - SentenceTransformers
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')  # Small but effective free model

# LLM for text generation - FLAN-T5-base (smaller free version)
tokenizer_llm = AutoTokenizer.from_pretrained("google/flan-t5-base")
model_llm = AutoModelForSeq2SeqLM.from_pretrained("google/flan-t5-base")

# QA model for direct answers
qa_tokenizer = AutoTokenizer.from_pretrained("deepset/roberta-base-squad2")
qa_model = AutoModelForQuestionAnswering.from_pretrained("TheBloke/Mistral-7B-Instruct-v0.2-GPTQ")
qa_pipeline = pipeline('question-answering', model=qa_model, tokenizer=qa_tokenizer)

# Vector Database
class VectorDatabase:
    def __init__(self, dim=384):  # MiniLM-L6 has 384 dimensions
        self.dimension = dim
        self.index = faiss.IndexFlatL2(dim)
        self.texts = []
        self.metadata = []
        self.loaded = False
    
    def add(self, texts, embeddings, metadata_list=None):
        if metadata_list is None:
            metadata_list = [{} for _ in texts]
        
        # Convert to numpy array if not already
        if isinstance(embeddings, torch.Tensor):
            embeddings = embeddings.numpy()
        
        # Normalize vectors for cosine similarity
        faiss.normalize_L2(embeddings)
        
        # Add to index
        self.index.add(embeddings)
        self.texts.extend(texts)
        self.metadata.extend(metadata_list)
    
    def search(self, query_embedding, k=5):
        # Ensure query embedding is properly formatted
        if isinstance(query_embedding, torch.Tensor):
            query_embedding = query_embedding.numpy()
            
        # Reshape if needed
        if len(query_embedding.shape) == 1:
            query_embedding = query_embedding.reshape(1, -1)
        
        # Normalize query vector
        faiss.normalize_L2(query_embedding)
        
        # Search
        distances, indices = self.index.search(query_embedding, k)
        
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(self.texts):  # Safety check
                results.append({
                    'text': self.texts[idx],
                    'score': float(distances[0][i]),
                    'metadata': self.metadata[idx]
                })
        
        return results
    
    def save(self, path):
        # Save the index
        faiss.write_index(self.index, os.path.join(path, "faiss.index"))
        
        # Save texts and metadata
        with open(os.path.join(path, "data.json"), 'w') as f:
            json.dump({'texts': self.texts, 'metadata': self.metadata}, f)
    
    def load(self, path):
        # Load the index
        self.index = faiss.read_index(os.path.join(path, "faiss.index"))
        
        # Load texts and metadata
        with open(os.path.join(path, "data.json"), 'r') as f:
            data = json.load(f)
            self.texts = data['texts']
            self.metadata = data['metadata']
        
        self.loaded = True
        return True

# Initialize vector database
vector_db = VectorDatabase()
try:
    vector_db.load(VECTOR_DB_PATH)
    print("Loaded existing vector database")
except:
    print("Creating new vector database")

# Document Processing Functions
def chunk_text(text, chunk_size=200, overlap=50):
    """Split text into overlapping chunks."""
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = []
    current_size = 0
    
    for sentence in sentences:
        sentence_tokens = len(sentence.split())
        
        if current_size + sentence_tokens > chunk_size and current_chunk:
            # Add current chunk to chunks
            chunks.append(' '.join(current_chunk))
            
            # Keep overlap sentences for next chunk
            overlap_size = 0
            overlap_chunk = []
            for s in reversed(current_chunk):
                s_tokens = len(s.split())
                if overlap_size + s_tokens <= overlap:
                    overlap_chunk.insert(0, s)
                    overlap_size += s_tokens
                else:
                    break
            
            current_chunk = overlap_chunk
            current_size = overlap_size
        
        current_chunk.append(sentence)
        current_size += sentence_tokens
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def process_pdf(file_path):
    """Extract text from PDF."""
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    
    return text

def process_docx(file_path):
    """Extract text from DOCX."""
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    return text

def process_csv(file_path):
    """Extract text from CSV by converting to descriptive text."""
    try:
        df = pd.read_csv(file_path)
        # Convert dataframe to readable text format
        text = f"CSV file with {len(df)} rows and {len(df.columns)} columns.\n"
        text += f"Columns: {', '.join(df.columns)}.\n\n"
        
        # Sample data summary
        text += "Sample data and statistics:\n"
        for col in df.columns:
            if pd.api.types.is_numeric_dtype(df[col]):
                text += f"Column {col}: min={df[col].min()}, max={df[col].max()}, "
                text += f"mean={df[col].mean():.2f}, median={df[col].median()}.\n"
            else:
                unique_vals = df[col].nunique()
                text += f"Column {col}: {unique_vals} unique values.\n"
                if unique_vals < 10:  # Only show all values if there are few
                    text += f"Values: {', '.join(df[col].unique().astype(str))}.\n"
                else:
                    text += f"Sample values: {', '.join(df[col].sample(5).astype(str))}.\n"
        
        # Add actual data rows as text
        text += "\nSample rows:\n"
        for i, row in df.head(5).iterrows():
            text += f"Row {i}: {' | '.join([f'{col}={val}' for col, val in row.items()])}.\n"
        
        return text
    except Exception as e:
        return f"Error processing CSV: {str(e)}"

def process_document(file_path):
    """Process any supported document type."""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return process_pdf(file_path)
    elif file_ext == '.docx':
        return process_docx(file_path)
    elif file_ext == '.csv':
        return process_csv(file_path)
    else:
        return f"Unsupported file type: {file_ext}"

def add_to_knowledge_base(file_path, document_id=None):
    """Process document and add to knowledge base."""
    if document_id is None:
        document_id = str(uuid.uuid4())
    
    # Extract text
    full_text = process_document(file_path)
    if full_text.startswith("Error") or full_text.startswith("Unsupported"):
        return {"status": "error", "message": full_text}
    
    # Chunk the text
    chunks = chunk_text(full_text)
    
    # Create metadata for each chunk
    metadata_list = [{"source": os.path.basename(file_path), "doc_id": document_id, "chunk_id": i} for i in range(len(chunks))]
    
    # Get embeddings for all chunks
    embeddings = embedding_model.encode(chunks)
    
    # Add to vector database
    vector_db.add(chunks, embeddings, metadata_list)
    
    # Save the updated database
    vector_db.save(VECTOR_DB_PATH)
    
    return {"status": "success", "message": f"Added {len(chunks)} chunks to knowledge base", "doc_id": document_id}

def build_prebuilt_knowledge_base():
    """Process all documents in the knowledge base folder."""
    for filename in os.listdir(KB_FOLDER):
        file_path = os.path.join(KB_FOLDER, filename)
        if os.path.isfile(file_path):
            print(f"Processing {filename} for prebuilt knowledge base...")
            result = add_to_knowledge_base(file_path)
            print(result)

# RAG Query Functions
def generate_answer(query, context_texts, max_length=100):
    """Generate an answer using the LLM based on the query and context."""
    # Prepare prompt with context
    prompt = f"""
    Based on the following information, please answer the question.
    
    Question: {query}
    
    Information:
    {' '.join(context_texts)}
    
    Answer:
    """
    
    inputs = tokenizer_llm(prompt, return_tensors="pt", max_length=512, truncation=True)
    
    # Generate
    outputs = model_llm.generate(
        inputs.input_ids,
        max_length=max_length,
        num_beams=4,
        early_stopping=True
    )
    
    # Decode
    answer = tokenizer_llm.decode(outputs[0], skip_special_tokens=True)
    
    return answer

def query_knowledge_base(query, top_k=5):
    """Query the knowledge base and return relevant chunks."""
    # Get query embedding
    query_embedding = embedding_model.encode([query])
    
    # Search vector database
    results = vector_db.search(query_embedding, k=top_k)
    
    # Extract texts
    context_texts = [r['text'] for r in results]
    sources = [r['metadata']['source'] for r in results]
    
    # Use QA model to get direct answer
    if len(context_texts) > 0:
        context = " ".join(context_texts)
        try:
            qa_result = qa_pipeline(question=query, context=context)
            answer = qa_result['answer']
            confidence = qa_result['score']
        except:
            # Fall back to LLM
            answer = generate_answer(query, context_texts)
            confidence = None
    else:
        answer = "I don't have enough information to answer that question."
        confidence = None
    
    return {
        "answer": answer,
        "confidence": confidence,
        "sources": list(set(sources)),  # Unique sources
        "contexts": context_texts
    }

# API Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/query', methods=['POST'])
def api_query():
    data = request.json
    query = data.get('query', '')
    
    if not query:
        return jsonify({"error": "No query provided"}), 400
    
    # Process query
    result = query_knowledge_base(query)
    
    return jsonify(result)

@app.route('/api/upload', methods=['POST'])
def api_upload():
    # Check if files were uploaded
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    # Save file temporarily
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)
    
    # Process file and add to knowledge base
    result = add_to_knowledge_base(file_path)
    
    # Clean up
    os.remove(file_path)
    
    return jsonify(result)

@app.route('/api/prebuilt-status', methods=['GET'])
def api_prebuilt_status():
    # Return status of prebuilt knowledge base
    status = {
        "total_chunks": len(vector_db.texts),
        "document_count": len(set([meta.get('doc_id') for meta in vector_db.metadata]))
    }
    return jsonify(status)

# HTML template
@app.route('/templates/index.html')
def get_index_template():
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Document RAG Chatbot</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 0; padding: 20px; }
            #chatbox { height: 300px; border: 1px solid #ccc; overflow-y: auto; padding: 10px; margin-bottom: 10px; }
            #inputArea { display: flex; margin-bottom: 20px; }
            #userInput { flex-grow: 1; padding: 8px; }
            button { padding: 8px 15px; background: #4CAF50; color: white; border: none; cursor: pointer; }
            #fileUpload { margin-bottom: 20px; }
            .message { margin-bottom: 10px; }
            .user { color: blue; }
            .bot { color: green; }
            .source { font-size: 0.8em; color: #666; margin-top: 5px; }
        </style>
    </head>
    <body>
        <h1>Document RAG Chatbot</h1>
        
        <div id="status">Loading knowledge base status...</div>
        
        <h2>Upload Documents</h2>
        <div id="fileUpload">
            <input type="file" id="documentFile" accept=".pdf,.docx,.csv">
            <button onclick="uploadDocument()">Upload</button>
            <div id="uploadStatus"></div>
        </div>
        
        <h2>Chat</h2>
        <div id="chatbox"></div>
        
        <div id="inputArea">
            <input type="text" id="userInput" placeholder="Ask a question...">
            <button onclick="sendMessage()">Send</button>
        </div>
        
        <script>
            // Load knowledge base status
            fetch('/api/prebuilt-status')
                .then(response => response.json())
                .then(data => {
                    document.getElementById('status').innerHTML = 
                        `Knowledge Base Status: ${data.total_chunks} chunks from ${data.document_count} documents`;
                });
            
            // Send message function
            function sendMessage() {
                const userInput = document.getElementById('userInput');
                const query = userInput.value.trim();
                
                if (query === '') return;
                
                // Add user message to chat
                addMessage('user', query);
                userInput.value = '';
                
                // Send query to backend
                fetch('/api/query', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json'
                    },
                    body: JSON.stringify({ query: query })
                })
                .then(response => response.json())
                .then(data => {
                    // Add bot response to chat
                    addMessage('bot', data.answer, data.sources);
                })
                .catch(error => {
                    console.error('Error:', error);
                    addMessage('bot', 'Sorry, there was an error processing your request.');
                });
            }
            
            // Add message to chat
            function addMessage(role, text, sources = []) {
                const chatbox = document.getElementById('chatbox');
                const messageDiv = document.createElement('div');
                messageDiv.className = `message ${role}`;
                
                const sender = role === 'user' ? 'You' : 'Bot';
                messageDiv.innerHTML = `<strong>${sender}:</strong> ${text}`;
                
                // Add sources if available
                if (sources && sources.length > 0) {
                    const sourceDiv = document.createElement('div');
                    sourceDiv.className = 'source';
                    sourceDiv.innerHTML = `Sources: ${sources.join(', ')}`;
                    messageDiv.appendChild(sourceDiv);
                }
                
                chatbox.appendChild(messageDiv);
                chatbox.scrollTop = chatbox.scrollHeight;
            }
            
            // Upload document function
            function uploadDocument() {
                const fileInput = document.getElementById('documentFile');
                const file = fileInput.files[0];
                
                if (!file) {
                    document.getElementById('uploadStatus').textContent = 'Please select a file first.';
                    return;
                }
                
                const formData = new FormData();
                formData.append('file', file);
                
                document.getElementById('uploadStatus').textContent = 'Uploading...';
                
                fetch('/api/upload', {
                    method: 'POST',
                    body: formData
                })
                .then(response => response.json())
                .then(data => {
                    document.getElementById('uploadStatus').textContent = data.message;
                    
                    // Refresh knowledge base status
                    return fetch('/api/prebuilt-status');
                })
                .then(response => response.json())
                .then(data => {
                    document.getElementById('status').innerHTML = 
                        `Knowledge Base Status: ${data.total_chunks} chunks from ${data.document_count} documents`;
                })
                .catch(error => {
                    console.error('Error:', error);
                    document.getElementById('uploadStatus').textContent = 'Error uploading file.';
                });
            }
            
            // Event listener for Enter key
            document.getElementById('userInput').addEventListener('keypress', function(e) {
                if (e.key === 'Enter') {
                    sendMessage();
                }
            });
            
            // Initial message
            addMessage('bot', 'Hello! I am your document RAG chatbot. You can ask me questions about the documents in my knowledge base or upload new documents for me to learn from.');
        </script>
    </body>
    </html>
    """

# Initialize prebuilt knowledge base
if not vector_db.loaded or len(vector_db.texts) == 0:
    build_prebuilt_knowledge_base()

# Start the app
if __name__ == '__main__':
    app.run(debug=True, port=5000)